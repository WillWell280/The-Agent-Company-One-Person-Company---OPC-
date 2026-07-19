# -*- coding: utf-8 -*-

import io
import unittest

from docx import Document
from openpyxl import Workbook

from studio.file_extractor import extract_file_content
from studio.generic_workflow import asset_summary


def make_pdf_with_text(text):
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
        ),
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    raw = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(raw))
        raw.extend(f"{index} 0 obj\n".encode("ascii"))
        raw.extend(obj)
        raw.extend(b"\nendobj\n")
    xref_offset = len(raw)
    raw.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    raw.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        raw.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    raw.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(raw)


class FileExtractorTests(unittest.TestCase):
    def assert_extracted(self, result, expected):
        self.assertEqual(result["extraction_status"], "success")
        self.assertIn(expected, result["extracted_text"])

    def test_extracts_txt_markdown_and_csv(self):
        self.assert_extracted(
            extract_file_content("brief.txt", "text/plain", "Objective: Launch a new product".encode("utf-8")),
            "Launch a new product",
        )
        self.assert_extracted(
            extract_file_content("notes.md", "text/markdown", b"# Project\n\nAcceptance criteria"),
            "Acceptance criteria",
        )
        csv_result = extract_file_content(
            "sales.csv",
            "text/csv",
            "Product,Revenue\nA,120\nB,90\n".encode("utf-8"),
        )
        self.assert_extracted(csv_result, "Product | Revenue")
        self.assertIn("A | 120", csv_result["extracted_text"])

    def test_extracts_pdf_text(self):
        result = extract_file_content("reference.pdf", "application/pdf", make_pdf_with_text("PDF Reference"))
        self.assert_extracted(result, "PDF Reference")
        self.assertIn("Page 1 text", result["extracted_text"])

    def test_extracts_docx_paragraphs_and_tables(self):
        document = Document()
        document.add_paragraph("Word requirement")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Item"
        table.cell(0, 1).text = "Owner"
        table.cell(1, 0).text = "Launch"
        table.cell(1, 1).text = "Alice"
        stream = io.BytesIO()
        document.save(stream)

        result = extract_file_content(
            "requirements.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            stream.getvalue(),
        )
        self.assert_extracted(result, "Word requirement")
        self.assertIn("Item | Owner", result["extracted_text"])

    def test_extracts_xlsx_sheets_and_rows(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Plan"
        sheet.append(["Task", "Status"])
        sheet.append(["Research", "Done"])
        stream = io.BytesIO()
        workbook.save(stream)
        workbook.close()

        result = extract_file_content(
            "plan.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            stream.getvalue(),
        )
        self.assert_extracted(result, "Worksheet: Plan")
        self.assertIn("Research | Done", result["extracted_text"])

    def test_images_remain_visual_and_extracted_text_enters_prompt(self):
        visual = extract_file_content("reference.png", "image/png", b"image-bytes")
        self.assertEqual(visual["extraction_status"], "visual")

        extracted = extract_file_content("brief.txt", "text/plain", b"Prompt-ready content")
        summary = asset_summary([
            {"name": "brief.txt", "mime": "text/plain", "size": 20, **extracted},
            {"name": "reference.png", "mime": "image/png", "size": 11, **visual},
        ])
        self.assertIn("Prompt-ready content", summary)
        self.assertIn("reference.png", summary)


if __name__ == "__main__":
    unittest.main()
