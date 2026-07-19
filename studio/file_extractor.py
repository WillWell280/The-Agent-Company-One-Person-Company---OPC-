# -*- coding: utf-8 -*-
"""Extract prompt-ready text and table summaries from uploaded documents."""

import csv
import io
import re
import zipfile
from pathlib import Path


MAX_EXTRACTED_CHARS = 24000
MAX_ARCHIVE_FILES = 5000
MAX_ARCHIVE_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
MAX_PDF_PAGES = 80
MAX_TABLES = 30
MAX_TABLE_ROWS = 200
MAX_TABLE_COLUMNS = 50
MAX_CELL_CHARS = 500


def _clean_cell(value):
    text = "" if value is None else str(value)
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) > MAX_CELL_CHARS:
        return text[:MAX_CELL_CHARS] + "..."
    return text


def _limit_text(text):
    text = str(text or "").replace("\x00", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if len(text) <= MAX_EXTRACTED_CHARS:
        return text, False
    return text[:MAX_EXTRACTED_CHARS].rstrip() + "\n\n[Content truncated because it exceeded the extraction limit.]", True


def _decode_text(raw):
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "big5", "latin-1"):
        try:
            return raw.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace"), "utf-8-replace"


def _safe_zip(raw):
    with zipfile.ZipFile(io.BytesIO(raw)) as archive:
        members = archive.infolist()
        if len(members) > MAX_ARCHIVE_FILES:
            raise ValueError("The archive contains more files than the safety limit allows.")
        if sum(item.file_size for item in members) > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
            raise ValueError("The uncompressed archive exceeds the 50 MB safety limit.")


def _format_table(rows, title):
    normalized = []
    for row in rows[:MAX_TABLE_ROWS]:
        cells = [_clean_cell(value) for value in list(row)[:MAX_TABLE_COLUMNS]]
        if any(cells):
            normalized.append(cells)
    if not normalized:
        return ""
    width = max(len(row) for row in normalized)
    lines = [title, f"Rows extracted: {len(normalized)}; columns: {width}"]
    lines.extend(" | ".join(row + [""] * (width - len(row))) for row in normalized)
    if len(rows) > MAX_TABLE_ROWS:
        lines.append(f"[Table truncated after {MAX_TABLE_ROWS} rows.]")
    return "\n".join(lines)


def _extract_plain_text(raw, format_name):
    text, encoding = _decode_text(raw)
    return f"Format: {format_name}\nEncoding: {encoding}\n\n{text}"


def _extract_csv(raw):
    text, encoding = _decode_text(raw)
    sample = text[:8192]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",\t;|")
    except csv.Error:
        dialect = csv.excel
    rows = []
    total_rows = 0
    for row in csv.reader(io.StringIO(text), dialect):
        total_rows += 1
        if len(rows) < MAX_TABLE_ROWS:
            rows.append(row)
    table = _format_table(rows, "CSV table content")
    if total_rows > MAX_TABLE_ROWS:
        table += f"\n[CSV truncated after {MAX_TABLE_ROWS} rows.]"
    return f"Format: CSV\nEncoding: {encoding}\nTotal rows: {total_rows}\n\n{table}"


def _extract_pdf(raw):
    import pdfplumber

    blocks = []
    table_count = 0
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        page_count = len(pdf.pages)
        for page_number, page in enumerate(pdf.pages[:MAX_PDF_PAGES], start=1):
            text = (page.extract_text() or "").strip()
            if text:
                blocks.append(f"## Page {page_number} text\n{text}")
            for table in page.extract_tables() or []:
                if table_count >= MAX_TABLES:
                    break
                table_count += 1
                rendered = _format_table(table, f"## Page {page_number}, table {table_count}")
                if rendered:
                    blocks.append(rendered)
        if page_count > MAX_PDF_PAGES:
            blocks.append(f"[The PDF has {page_count} pages. Only the first {MAX_PDF_PAGES} were extracted.]")
    return f"Format: PDF\nPages: {page_count}\nTables detected: {table_count}\n\n" + "\n\n".join(blocks)


def _extract_docx(raw):
    from docx import Document

    _safe_zip(raw)
    document = Document(io.BytesIO(raw))
    blocks = []
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    if paragraphs:
        blocks.append("## Body\n" + "\n".join(paragraphs))
    for index, table in enumerate(document.tables[:MAX_TABLES], start=1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        rendered = _format_table(rows, f"## Table {index}")
        if rendered:
            blocks.append(rendered)
    return (
        f"Format: Word DOCX\nBody paragraphs: {len(paragraphs)}\nTables: {len(document.tables)}\n\n"
        + "\n\n".join(blocks)
    )


def _extract_xlsx(raw):
    from openpyxl import load_workbook

    _safe_zip(raw)
    workbook = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    blocks = []
    try:
        for sheet in workbook.worksheets:
            rows = []
            total_rows = 0
            for row in sheet.iter_rows(values_only=True):
                total_rows += 1
                if len(rows) < MAX_TABLE_ROWS:
                    rows.append(list(row))
            rendered = _format_table(rows, f"## Worksheet: {sheet.title}")
            if rendered:
                blocks.append(f"Total rows: {total_rows}\n{rendered}")
    finally:
        workbook.close()
    return f"Format: Excel XLSX\nWorksheets: {len(workbook.sheetnames)}\n\n" + "\n\n".join(blocks)


def _extract_xls(raw):
    import xlrd

    workbook = xlrd.open_workbook(file_contents=raw, on_demand=True)
    blocks = []
    try:
        for sheet in workbook.sheets():
            rows = [
                [sheet.cell_value(row_index, column_index) for column_index in range(min(sheet.ncols, MAX_TABLE_COLUMNS))]
                for row_index in range(min(sheet.nrows, MAX_TABLE_ROWS))
            ]
            rendered = _format_table(rows, f"## Worksheet: {sheet.name}")
            if rendered:
                blocks.append(f"Total rows: {sheet.nrows}\n{rendered}")
    finally:
        workbook.release_resources()
    return f"Format: Excel XLS\nWorksheets: {workbook.nsheets}\n\n" + "\n\n".join(blocks)


def extract_file_content(filename, mime, raw):
    """Return extraction metadata. Original bytes are never changed."""
    name = str(filename or "Untitled attachment")
    suffix = Path(name).suffix.lower()
    mime = str(mime or "").lower()
    if mime.startswith("image/"):
        return {
            "extraction_status": "visual",
            "extraction_format": "image",
            "extraction_note": "This image will be passed as visual context to vision-capable models.",
        }

    handlers = {
        ".txt": lambda: _extract_plain_text(raw, "TXT"),
        ".md": lambda: _extract_plain_text(raw, "Markdown"),
        ".markdown": lambda: _extract_plain_text(raw, "Markdown"),
        ".csv": lambda: _extract_csv(raw),
        ".pdf": lambda: _extract_pdf(raw),
        ".docx": lambda: _extract_docx(raw),
        ".xlsx": lambda: _extract_xlsx(raw),
        ".xlsm": lambda: _extract_xlsx(raw),
        ".xls": lambda: _extract_xls(raw),
    }
    handler = handlers.get(suffix)
    if handler is None:
        return {
            "extraction_status": "unsupported",
            "extraction_format": suffix.lstrip(".") or mime or "unknown",
            "extraction_note": "Text extraction is not available for this format. The original attachment is still preserved.",
        }

    try:
        extracted, truncated = _limit_text(handler())
        if not extracted:
            return {
                "extraction_status": "empty",
                "extraction_format": suffix.lstrip("."),
                "extraction_note": "The file was readable, but no usable text or tables were found.",
            }
        return {
            "extraction_status": "success",
            "extraction_format": suffix.lstrip("."),
            "extracted_text": extracted,
            "extracted_chars": len(extracted),
            "extraction_truncated": truncated,
            "extraction_note": "Text and table summaries were extracted and will be included in the task context.",
        }
    except Exception as exc:
        return {
            "extraction_status": "error",
            "extraction_format": suffix.lstrip("."),
            "extraction_note": f"Content extraction failed: {exc.__class__.__name__}: {exc}",
        }
